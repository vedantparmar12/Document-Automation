"""
Multi-format documentation exporter with theme support.

This module provides comprehensive export capabilities for documentation
in various formats including HTML, PDF, Word, Confluence, and more.
"""

import json
import logging
import tempfile
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)


class ExportFormat(Enum):
    """Supported export formats."""
    HTML = "html"
    PDF = "pdf" 
    DOCX = "docx"
    MARKDOWN = "markdown"
    CONFLUENCE = "confluence"
    NOTION = "notion"
    JSON = "json"
    EPUB = "epub"


class Theme(Enum):
    """Available themes."""
    DEFAULT = "default"
    DARK = "dark"
    MINIMAL = "minimal"
    CORPORATE = "corporate"
    GITHUB = "github"
    MATERIAL = "material"


@dataclass
class ExportOptions:
    """Export configuration options."""
    format: ExportFormat
    theme: Theme = Theme.DEFAULT
    include_toc: bool = True
    include_diagrams: bool = True
    include_code_highlighting: bool = True
    include_search: bool = True
    custom_css: Optional[str] = None
    output_path: Optional[str] = None
    title: Optional[str] = None


class DocumentationExporter:
    """
    Export documentation to multiple formats with theme support.
    """
    
    def __init__(self):
        self.themes = self._load_themes()
        self.templates = self._load_templates()
        
    def export_documentation(
        self,
        analysis_data: Dict[str, Any],
        options: ExportOptions
    ) -> Dict[str, Any]:
        """
        Export documentation in the specified format.
        
        Args:
            analysis_data: Complete analysis data
            options: Export configuration
            
        Returns:
            Export result with file path and metadata
        """
        try:
            logger.info(f"Exporting documentation to {options.format.value} with {options.theme.value} theme")
            
            if options.format == ExportFormat.HTML:
                return self._export_html(analysis_data, options)
            elif options.format == ExportFormat.PDF:
                return self._export_pdf(analysis_data, options)
            elif options.format == ExportFormat.DOCX:
                return self._export_docx(analysis_data, options)
            elif options.format == ExportFormat.MARKDOWN:
                return self._export_markdown(analysis_data, options)
            elif options.format == ExportFormat.CONFLUENCE:
                return self._export_confluence(analysis_data, options)
            elif options.format == ExportFormat.NOTION:
                return self._export_notion(analysis_data, options)
            elif options.format == ExportFormat.JSON:
                return self._export_json(analysis_data, options)
            elif options.format == ExportFormat.EPUB:
                return self._export_epub(analysis_data, options)
            else:
                raise ValueError(f"Unsupported export format: {options.format}")
                
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'format': options.format.value
            }
    
    def export_multiple_formats(
        self,
        analysis_data: Dict[str, Any],
        formats: List[ExportFormat],
        theme: Theme = Theme.DEFAULT,
        output_dir: Optional[str] = None
    ) -> Dict[str, Dict[str, Any]]:
        """Export documentation in multiple formats simultaneously."""
        results = {}
        
        for format_type in formats:
            options = ExportOptions(
                format=format_type,
                theme=theme,
                output_path=output_dir
            )
            
            result = self.export_documentation(analysis_data, options)
            results[format_type.value] = result
            
        return results
    
    def _export_html(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export to HTML format."""
        try:
            from src.generators.interactive_doc_generator import InteractiveDocumentationGenerator
            
            generator = InteractiveDocumentationGenerator()
            html_content = generator.generate_interactive_docs(
                analysis_data=analysis_data,
                title=options.title or "Project Documentation",
                theme=options.theme.value,
                include_search=options.include_search,
                include_navigation=True,
                include_live_diagrams=options.include_diagrams
            )
            
            # Apply custom CSS if provided
            if options.custom_css:
                html_content = self._apply_custom_css(html_content, options.custom_css)
            
            # Save to file
            output_path = self._get_output_path(options, "documentation.html")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'format': 'html',
                'output_path': output_path,
                'file_size': len(html_content.encode('utf-8')),
                'features': {
                    'interactive': True,
                    'searchable': options.include_search,
                    'responsive': True,
                    'diagrams': options.include_diagrams
                }
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'html'}
    
    def _export_pdf(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export to PDF format."""
        try:
            # First generate HTML
            html_options = ExportOptions(
                format=ExportFormat.HTML,
                theme=options.theme,
                include_toc=options.include_toc,
                include_diagrams=options.include_diagrams,
                include_code_highlighting=options.include_code_highlighting,
                include_search=False,  # Not relevant for PDF
                title=options.title
            )
            
            html_result = self._export_html(analysis_data, html_options)
            
            if not html_result.get('success'):
                return html_result
            
            # Convert HTML to PDF using weasyprint
            try:
                from weasyprint import HTML, CSS
                
                html_path = html_result['output_path']
                pdf_path = self._get_output_path(options, "documentation.pdf")
                
                # Add PDF-specific CSS
                pdf_css = self._get_pdf_css(options.theme)
                
                html_doc = HTML(filename=html_path)
                css_doc = CSS(string=pdf_css)
                
                html_doc.write_pdf(pdf_path, stylesheets=[css_doc])
                
                return {
                    'success': True,
                    'format': 'pdf',
                    'output_path': pdf_path,
                    'file_size': Path(pdf_path).stat().st_size,
                    'features': {
                        'printable': True,
                        'bookmarks': options.include_toc,
                        'diagrams': options.include_diagrams
                    }
                }
                
            except ImportError:
                return {
                    'success': False,
                    'error': 'weasyprint not installed. Install with: pip install weasyprint',
                    'format': 'pdf'
                }
                
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'pdf'}
    
    def _export_docx(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export to Word document format."""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Set document title
            title = options.title or "Project Documentation"
            doc_title = doc.add_heading(title, 0)
            
            # Add sections based on analysis data
            self._add_docx_sections(doc, analysis_data, options)
            
            # Save document
            output_path = self._get_output_path(options, "documentation.docx")
            doc.save(output_path)
            
            return {
                'success': True,
                'format': 'docx',
                'output_path': output_path,
                'file_size': Path(output_path).stat().st_size,
                'features': {
                    'editable': True,
                    'formatted': True,
                    'compatible': True
                }
            }
            
        except ImportError:
            return {
                'success': False,
                'error': 'python-docx not installed. Install with: pip install python-docx',
                'format': 'docx'
            }
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'docx'}
    
    def _export_markdown(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export to Markdown format."""
        try:
            markdown_content = self._generate_markdown_content(analysis_data, options)
            
            output_path = self._get_output_path(options, "documentation.md")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return {
                'success': True,
                'format': 'markdown',
                'output_path': output_path,
                'file_size': len(markdown_content.encode('utf-8')),
                'features': {
                    'portable': True,
                    'version_controllable': True,
                    'readable': True,
                    'diagrams_supported': options.include_diagrams
                }
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'markdown'}
    
    def _export_confluence(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export to Confluence wiki format."""
        try:
            confluence_content = self._generate_confluence_content(analysis_data, options)
            
            output_path = self._get_output_path(options, "documentation_confluence.txt")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(confluence_content)
            
            return {
                'success': True,
                'format': 'confluence',
                'output_path': output_path,
                'file_size': len(confluence_content.encode('utf-8')),
                'features': {
                    'wiki_formatted': True,
                    'collaborative': True,
                    'macros_supported': True
                },
                'instructions': 'Copy content to Confluence page editor'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'confluence'}
    
    def _export_notion(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export to Notion format."""
        try:
            notion_blocks = self._generate_notion_blocks(analysis_data, options)
            
            output_path = self._get_output_path(options, "notion_blocks.json")
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(notion_blocks, f, indent=2)
            
            return {
                'success': True,
                'format': 'notion',
                'output_path': output_path,
                'file_size': Path(output_path).stat().st_size,
                'features': {
                    'block_structured': True,
                    'interactive': True,
                    'embeddable': True
                },
                'instructions': 'Import blocks using Notion API'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'notion'}
    
    def _export_json(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export raw data as JSON."""
        try:
            # Enhanced JSON with export metadata
            export_data = {
                'metadata': {
                    'export_format': 'json',
                    'theme': options.theme.value,
                    'generated_at': self._get_timestamp(),
                    'title': options.title or "Project Documentation"
                },
                'analysis_data': analysis_data
            }
            
            output_path = self._get_output_path(options, "documentation.json")
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, default=str)
            
            return {
                'success': True,
                'format': 'json',
                'output_path': output_path,
                'file_size': Path(output_path).stat().st_size,
                'features': {
                    'machine_readable': True,
                    'structured': True,
                    'processable': True
                }
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'json'}
    
    def _export_epub(self, analysis_data: Dict[str, Any], options: ExportOptions) -> Dict[str, Any]:
        """Export to EPUB format."""
        try:
            from ebooklib import epub
            
            book = epub.EpubBook()
            
            # Set metadata
            title = options.title or "Project Documentation"
            book.set_identifier('documentation')
            book.set_title(title)
            book.set_language('en')
            book.add_author('MCP Document Automation')
            
            # Generate chapters
            chapters = self._generate_epub_chapters(analysis_data, options)
            
            for chapter in chapters:
                book.add_item(chapter)
            
            # Define Table of Contents
            book.toc = [(epub.Section('Documentation'), chapters)]
            
            # Add navigation
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Create spine
            book.spine = ['nav'] + chapters
            
            # Save EPUB
            output_path = self._get_output_path(options, "documentation.epub")
            epub.write_epub(output_path, book, {})
            
            return {
                'success': True,
                'format': 'epub',
                'output_path': output_path,
                'file_size': Path(output_path).stat().st_size,
                'features': {
                    'portable': True,
                    'reflowable': True,
                    'accessible': True
                }
            }
            
        except ImportError:
            return {
                'success': False,
                'error': 'EbookLib not installed. Install with: pip install EbookLib',
                'format': 'epub'
            }
        except Exception as e:
            return {'success': False, 'error': str(e), 'format': 'epub'}
    
    def _generate_markdown_content(self, analysis_data: Dict[str, Any], options: ExportOptions) -> str:
        """Generate markdown content from analysis data."""
        content = []
        
        # Title
        title = options.title or "Project Documentation"
        content.append(f"# {title}\n")
        
        # Table of contents
        if options.include_toc:
            content.append("## Table of Contents\n")
            content.append("- [Project Overview](#project-overview)")
            if 'technology_stack' in analysis_data:
                content.append("- [Technology Stack](#technology-stack)")
            if 'file_structure' in analysis_data:
                content.append("- [Project Structure](#project-structure)")
            if analysis_data.get('api_endpoints'):
                content.append("- [API Endpoints](#api-endpoints)")
            if analysis_data.get('database_schemas'):
                content.append("- [Database Schemas](#database-schemas)")
            content.append("")
        
        # Overview section
        content.append("## Project Overview\n")
        content.append("This documentation was automatically generated from codebase analysis.\n")
        
        if 'total_files' in analysis_data:
            content.append(f"- **Total Files**: {analysis_data['total_files']}")
        if 'analyzed_files' in analysis_data:
            content.append(f"- **Analyzed Files**: {analysis_data['analyzed_files']}")
        content.append("")
        
        # Technology Stack
        if 'technology_stack' in analysis_data:
            content.append("## Technology Stack\n")
            tech_stack = analysis_data['technology_stack']
            
            for category, items in tech_stack.items():
                if items and category != 'confidence_scores':
                    content.append(f"### {category.replace('_', ' ').title()}")
                    if isinstance(items, list):
                        for item in items:
                            content.append(f"- {item}")
                    else:
                        content.append(f"- {items}")
                    content.append("")
        
        # File Structure
        if 'file_structure' in analysis_data:
            content.append("## Project Structure\n")
            structure = analysis_data['file_structure']
            content.append(f"**Statistics:**")
            content.append(f"- Total Lines: {structure.get('total_lines', 0):,}")
            content.append(f"- File Types: {len(structure.get('file_types', {}))}")
            content.append("")
            
            if structure.get('file_types'):
                content.append("**File Types:**")
                for ext, count in structure['file_types'].items():
                    content.append(f"- `{ext or 'no extension'}`: {count} files")
                content.append("")
        
        # API Endpoints
        if analysis_data.get('api_endpoints'):
            content.append("## API Endpoints\n")
            endpoints = analysis_data['api_endpoints']
            
            for endpoint in endpoints[:20]:  # Limit display
                methods = ', '.join(endpoint.get('methods', ['GET']))
                content.append(f"### `{methods}` {endpoint.get('path', '')}")
                content.append(f"- **File**: `{endpoint.get('file', '')}`")
                content.append(f"- **Framework**: {endpoint.get('framework', 'unknown')}")
                content.append("")
        
        # Database Schemas
        if analysis_data.get('database_schemas'):
            content.append("## Database Schemas\n")
            schemas = analysis_data['database_schemas']
            
            for i, schema in enumerate(schemas, 1):
                content.append(f"### Schema {i}")
                content.append(f"- **Type**: {schema.get('schema_type', 'Unknown')}")
                content.append(f"- **Tables**: {len(schema.get('tables', []))}")
                
                if schema.get('tables'):
                    content.append("**Tables:**")
                    for table in schema['tables'][:10]:
                        content.append(f"- `{table}`")
                    if len(schema['tables']) > 10:
                        content.append(f"- ... and {len(schema['tables']) - 10} more")
                content.append("")
        
        # Dependencies
        if analysis_data.get('dependencies'):
            deps = analysis_data['dependencies']
            if deps.get('dependencies'):
                content.append("## Dependencies\n")
                for dep, version in list(deps['dependencies'].items())[:15]:
                    content.append(f"- `{dep}`: {version}")
                if len(deps['dependencies']) > 15:
                    content.append(f"- ... and {len(deps['dependencies']) - 15} more")
                content.append("")
        
        # Mermaid diagrams
        if options.include_diagrams and analysis_data:
            content.append("## Architecture Diagram\n")
            content.append("```mermaid")
            content.append("graph TD")
            content.append("    A[Project Root] --> B[Source Code]")
            content.append("    A --> C[Configuration]")
            content.append("    A --> D[Documentation]")
            content.append("```\n")
        
        return '\n'.join(content)
    
    def _generate_confluence_content(self, analysis_data: Dict[str, Any], options: ExportOptions) -> str:
        """Generate Confluence wiki markup."""
        content = []
        
        title = options.title or "Project Documentation"
        content.append(f"h1. {title}\n")
        
        content.append("This documentation was automatically generated from codebase analysis.\n")
        
        if 'technology_stack' in analysis_data:
            content.append("h2. Technology Stack\n")
            tech_stack = analysis_data['technology_stack']
            
            for category, items in tech_stack.items():
                if items and category != 'confidence_scores':
                    content.append(f"h3. {category.replace('_', ' ').title()}")
                    if isinstance(items, list):
                        for item in items:
                            content.append(f"* {item}")
                    else:
                        content.append(f"* {items}")
                    content.append("")
        
        return '\n'.join(content)
    
    def _generate_notion_blocks(self, analysis_data: Dict[str, Any], options: ExportOptions) -> List[Dict[str, Any]]:
        """Generate Notion blocks structure."""
        blocks = []
        
        # Title block
        blocks.append({
            "type": "heading_1",
            "heading_1": {
                "rich_text": [{"type": "text", "text": {"content": options.title or "Project Documentation"}}]
            }
        })
        
        # Overview paragraph
        blocks.append({
            "type": "paragraph",
            "paragraph": {
                "rich_text": [{"type": "text", "text": {"content": "This documentation was automatically generated from codebase analysis."}}]
            }
        })
        
        return blocks
    
    def _generate_epub_chapters(self, analysis_data: Dict[str, Any], options: ExportOptions) -> List:
        """Generate EPUB chapters."""
        from ebooklib import epub
        
        chapters = []
        
        # Overview chapter
        chapter1 = epub.EpubHtml(title='Overview', file_name='overview.xhtml', lang='en')
        chapter1.content = f'''
        <html>
        <head><title>Overview</title></head>
        <body>
            <h1>Project Overview</h1>
            <p>This documentation was automatically generated from codebase analysis.</p>
        </body>
        </html>'''
        chapters.append(chapter1)
        
        return chapters
    
    def _add_docx_sections(self, doc, analysis_data: Dict[str, Any], options: ExportOptions) -> None:
        """Add sections to Word document."""
        # Overview section
        doc.add_heading('Project Overview', level=1)
        doc.add_paragraph('This documentation was automatically generated from codebase analysis.')
        
        # Add more sections based on analysis_data...
        if 'technology_stack' in analysis_data:
            doc.add_heading('Technology Stack', level=1)
            tech_stack = analysis_data['technology_stack']
            
            for category, items in tech_stack.items():
                if items and category != 'confidence_scores':
                    doc.add_heading(category.replace('_', ' ').title(), level=2)
                    if isinstance(items, list):
                        for item in items:
                            p = doc.add_paragraph()
                            p.add_run(f"- {item}")
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"- {items}")
    
    def _get_output_path(self, options: ExportOptions, default_filename: str) -> str:
        """Get output file path."""
        if options.output_path:
            if Path(options.output_path).is_dir():
                return str(Path(options.output_path) / default_filename)
            else:
                return options.output_path
        else:
            return default_filename
    
    def _get_timestamp(self) -> str:
        """Get current timestamp."""
        from datetime import datetime
        return datetime.now().isoformat()
    
    def _apply_custom_css(self, html_content: str, custom_css: str) -> str:
        """Apply custom CSS to HTML content."""
        # Find closing </style> tag and insert custom CSS before it
        style_end = html_content.find('</style>')
        if style_end != -1:
            return html_content[:style_end] + custom_css + html_content[style_end:]
        return html_content
    
    def _get_pdf_css(self, theme: Theme) -> str:
        """Get PDF-specific CSS."""
        return '''
        @page {
            margin: 1in;
            @bottom-center {
                content: counter(page);
            }
        }
        
        body {
            font-size: 12pt;
            line-height: 1.4;
        }
        
        .header {
            display: none;
        }
        
        .sidebar {
            display: none;
        }
        
        .main-content {
            margin: 0;
            padding: 0;
        }
        '''
    
    def _load_themes(self) -> Dict[str, Dict[str, str]]:
        """Load theme configurations."""
        return {
            'default': {'primary_color': '#007bff', 'background': '#ffffff'},
            'dark': {'primary_color': '#0d6efd', 'background': '#1a1a1a'},
            'minimal': {'primary_color': '#666666', 'background': '#ffffff'},
            'corporate': {'primary_color': '#2c3e50', 'background': '#f8f9fa'},
            'github': {'primary_color': '#0969da', 'background': '#ffffff'},
            'material': {'primary_color': '#1976d2', 'background': '#fafafa'}
        }
    
    def _load_templates(self) -> Dict[str, str]:
        """Load export templates."""
        return {
            'html': 'HTML template',
            'markdown': 'Markdown template',
            'confluence': 'Confluence template'
        }
